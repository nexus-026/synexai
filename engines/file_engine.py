import os
import io
from typing import Dict, Any
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


class FileEngine:
    async def analyze(self, message: str, file_path: str | None = None, file_bytes: bytes | None = None,
                      filename: str = "upload") -> Dict[str, Any]:
        if not file_path and not file_bytes:
            return {
                "success": False,
                "type": "analysis",
                "response": "No file provided for analysis.",
                "sources": [],
            }

        ext = Path(filename).suffix.lower()
        content = ""

        if ext == ".pdf":
            content = self._read_pdf(file_path, file_bytes)
        elif ext in (".docx", ".doc"):
            content = self._read_docx(file_path, file_bytes)
        elif ext in (".txt", ".md", ".csv", ".py", ".js", ".html", ".css", ".php", ".sql"):
            content = self._read_text(file_path, file_bytes)
        elif ext in (".xlsx", ".xls"):
            content = self._read_excel(file_path, file_bytes)
        else:
            content = self._read_text(file_path, file_bytes) or "[Binary or unsupported file type]"

        # Summarize / analyze
        summary = content[:2000] + ("..." if len(content) > 2000 else "")
        analysis = self._analyze_content(content, ext)

        return {
            "success": True,
            "type": "analysis",
            "response": f"**File Analysis: {filename}**\n\n**Summary:**\n{summary}\n\n**Analysis:**\n{analysis}",
            "file_info": {
                "filename": filename,
                "size": len(file_bytes) if file_bytes else os.path.getsize(file_path) if file_path else 0,
                "extension": ext,
            },
            "sources": [],
        }

    def _read_pdf(self, path: str | None, data: bytes | None) -> str:
        if not PdfReader:
            return "[PyPDF2 not installed]"
        try:
            if data:
                reader = PdfReader(io.BytesIO(data))
            else:
                reader = PdfReader(path)
            return "\n".join([p.extract_text() or "" for p in reader.pages])
        except Exception as e:
            return f"[PDF read error: {e}]"

    def _read_docx(self, path: str | None, data: bytes | None) -> str:
        if not Document:
            return "[python-docx not installed]"
        try:
            if data:
                doc = Document(io.BytesIO(data))
            else:
                doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"[DOCX read error: {e}]"

    def _read_text(self, path: str | None, data: bytes | None) -> str:
        try:
            if data:
                return data.decode("utf-8", errors="ignore")
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"[Text read error: {e}]"

    def _read_excel(self, path: str | None, data: bytes | None) -> str:
        if not load_workbook:
            return "[openpyxl not installed]"
        try:
            if data:
                wb = load_workbook(io.BytesIO(data))
            else:
                wb = load_workbook(path)
            lines = []
            for sheet in wb.worksheets[:3]:
                lines.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    lines.append(" | ".join(str(c) if c is not None else "" for c in row))
            return "\n".join(lines)
        except Exception as e:
            return f"[Excel read error: {e}]"

    def _analyze_content(self, content: str, ext: str) -> str:
        lines = content.split("\n")
        words = content.split()
        analysis = f"- Total lines: {len(lines)}\n- Total words: {len(words)}\n"
        if ext in (".py", ".js", ".php", ".java", ".cpp", ".c", ".ts"):
            funcs = len([l for l in lines if l.strip().startswith(("def ", "function ", "public ", "private "))])
            analysis += f"- Approximate functions/classes: {funcs}\n"
            if "TODO" in content or "FIXME" in content:
                analysis += "- ⚠️ Contains TODO/FIXME markers.\n"
        if ext in (".html", ".css"):
            tags = len(re.findall(r'<[a-zA-Z][^>]*>', content))
            analysis += f"- Approximate HTML tags: {tags}\n"
        if ext == ".csv":
            rows = len(lines)
            analysis += f"- CSV rows (incl. header): {rows}\n"
        return analysis


import re
